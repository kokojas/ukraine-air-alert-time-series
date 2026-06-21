"""DOCX report generation."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches

from air_alerts.constants import DISCLAIMER, PROJECT_TITLE, SOURCE_LEDGER


def write_report(analysis: dict[str, object], output_path: Path, chart_paths: dict[str, Path]) -> Path:
    document = Document()
    meta = analysis["meta"]
    summary = analysis["summary_last12"]
    predictions = analysis["predictions"]
    backtest = analysis["backtest"]

    document.add_heading(PROJECT_TITLE, level=0)
    document.add_paragraph(
        "A reproducible Python mini-project for analyzing the burden of air raid alerts in Ukraine. "
        f"The main window is {meta['last12_start']} to {meta['latest_date']}, with all available history used as context."
    )
    document.add_paragraph(DISCLAIMER)

    document.add_heading("Key Findings", level=1)
    if not summary.empty:
        top = summary.iloc[0]
        document.add_paragraph(
            f"The highest total alert burden in the last 12 months is {top['oblast']} "
            f"with {top['total_alert_hours']:.1f} alert hours."
        )
        night = summary.sort_values("total_night_hours", ascending=False).iloc[0]
        document.add_paragraph(
            f"The largest nighttime burden is {night['oblast']} with {night['total_night_hours']:.1f} night alert hours."
        )
    document.add_paragraph(
        "The baseline model is intentionally simple: recent 7-day mean, 30-day mean, and same-weekday history. "
        "Its purpose is to create a benchmark for validation, not a safety product."
    )

    _add_picture(document, chart_paths.get("burden"), "Figure 1. Top oblasts by alert burden.")
    _add_picture(document, chart_paths.get("focus"), "Figure 2. Focus regions, rolling 30-day burden.")
    _add_picture(document, chart_paths.get("corr"), "Figure 3. Co-alert similarity among focus regions.")
    _add_picture(document, chart_paths.get("pipeline"), "Figure 4. Reproducible analysis pipeline.")

    document.add_heading("Top Regions", level=1)
    _add_table(
        document,
        summary.head(10),
        ["oblast", "total_alert_hours", "total_night_hours", "alert_count", "night_share"],
    )

    document.add_heading("Baseline Predictions", level=1)
    _add_table(document, predictions.head(10), ["oblast", "target_date", "predicted_alert_hours", "model"])

    document.add_heading("Backtest", level=1)
    _add_table(document, backtest.head(10), ["oblast", "mae_hours", "n_test_days"])

    document.add_heading("Source and Method Ledger", level=1)
    ledger = pd.DataFrame(SOURCE_LEDGER)
    _add_table(document, ledger, ["layer", "source", "status", "caveat"])

    document.add_heading("Limitations", level=1)
    limitations = [
        "Air raid alerts are public-warning events, not direct records of attacks.",
        "From late 2025, official records become more granular in some places, so oblast aggregation is conservative.",
        "Night burden is based on time overlap with a fixed local-night window.",
        "The baseline model is not a life-safety decision tool.",
    ]
    for item in limitations:
        document.add_paragraph(item, style="List Bullet")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _add_picture(document: Document, path: Path | None, caption: str) -> None:
    if not path or not path.exists():
        return
    document.add_picture(str(path), width=Inches(6.5))
    document.add_paragraph(caption)


def _add_table(document: Document, frame: pd.DataFrame, columns: list[str]) -> None:
    visible_columns = [column for column in columns if column in frame.columns]
    if frame.empty or not visible_columns:
        document.add_paragraph("No data available.")
        return
    table = document.add_table(rows=1, cols=len(visible_columns))
    table.style = "Table Grid"
    for i, column in enumerate(visible_columns):
        table.rows[0].cells[i].text = column
    for row in frame[visible_columns].itertuples(index=False):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = _format_value(value)


def _format_value(value: object) -> str:
    if isinstance(value, float):
        return f"{value:.3f}".rstrip("0").rstrip(".")
    return str(value)
